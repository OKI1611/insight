# -*- coding: utf-8 -*-
"""정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션 · 한영대역 — Word(.docx) 생성

  PDF 정본(tools/build_pdfs.py build_parallel)과 같은 기준(2026-08-03 확정):
    신국판 152×225 · 존더반 병렬 표준형(좌 한글/우 영어 절 정렬) ·
    큰 숫자 장 표기 · 소제목 수록(표 사이 전체 폭) · 단 구분선(표 안쪽 세로선) ·
    한글 따옴표류 제거 · 쪽번호(바닥글)

  ⚠ OOXML 주의(과거 손상 사고): tblPr 자식 순서를 지켜야 워드가 파일을 연다.
    - tblW/tblLayout 은 python-docx API(autofit·width)로만 만지고 직접 append 금지
    - tblBorders 는 insert_element_before 로 스키마 위치에 삽입
  실행: python tools/build_parallel_docx.py
"""
import json, io, os, re, sys, time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
books = json.load(io.open(os.path.join(ROOT, "bible", "books.json"), encoding="utf-8"))

# 인수로 앞 N권만 생성(무결성 스모크 테스트용). 없으면 전권 정본.
LIMIT = int(sys.argv[1]) if len(sys.argv) > 1 else 0
if LIMIT:
    books = books[:LIMIT]
    OUT = os.path.join(ROOT, "책원고", "출판준비", "_smoke_parallel.docx")
else:
    OUT = os.path.join(ROOT, "책원고", "정본역킹제임스성경_한영대역.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

try:
    HEADINGS = json.load(io.open(os.path.join(ROOT, "bible", "headings", "ko.json"), encoding="utf-8"))
except Exception:
    HEADINGS = {}
def headings_for(book, ch):
    return {v: t for v, t in HEADINGS.get(book, {}).get(str(ch), [])}

_KO_QUOTES = re.compile(u'[“”‘’"\'「」『』《》〈〉]')
def clean_ko(s):
    return re.sub(r"\s{2,}", " ", _KO_QUOTES.sub("", str(s))).strip()

def en_ch(b, c):
    p = os.path.join(ROOT, "bible", "en", "%s-%d.json" % (b, c))
    return json.loads(io.open(p, encoding="utf-8").read()) if os.path.exists(p) else None

FONT  = "Noto Serif CJK KR"
GREEN = RGBColor(0x00, 0x59, 0x3c)
GRAY  = RGBColor(0x8a, 0x8a, 0x8a)
INK   = RGBColor(0x23, 0x2d, 0x28)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(15.2), Cm(22.5)        # 신국판
sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.4)
sec.left_margin = sec.right_margin = Cm(1.2)

st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(8.8)
st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.4

def kf(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if size: run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color
    return run

def center(txt, size, bold=False, color=None, before=0, after=6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    kf(p.add_run(txt), size, bold, color)
    return p

def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "17"); rPr.append(sz)
    r.append(rPr); t = OxmlElement("w:t"); t.text = "1"; r.append(t)
    fld.append(r); p._p.append(fld)

def set_insideV(tbl):
    """표 안쪽 세로 구분선 — tblBorders 를 스키마 위치에 삽입(순서 위반 금지)"""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    iv = OxmlElement("w:insideV")
    iv.set(qn("w:val"), "single"); iv.set(qn("w:sz"), "4")
    iv.set(qn("w:space"), "0"); iv.set(qn("w:color"), "C9C2B4")
    borders.append(iv)
    tblPr.insert_element_before(borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")

# ── 표지 (1단) ──
for _ in range(5): doc.add_paragraph()
# 역본명이 27자라 한 줄에 안 들어감 → 2줄로 나눠 배치
center("정본역(正本譯) 킹제임스 성경", 24, True, INK, 0, 2)
center("헤리티지 에디션", 16, True, INK)
center("한영대역", 19, True, GREEN)
center("THE HERITAGE KJV · AUTHENTIC VERSION · 1611", 10, False, GRAY)
for _ in range(8): doc.add_paragraph()
center("바이블 인사이트 출판사", 12, True)

# ── 판권 ──
doc.add_page_break()
for _ in range(1): doc.add_paragraph()
def line(txt, size=9, bold=False, color=None, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    kf(p.add_run(txt), size, bold, color)
line("일러두기 · 판권", 13, True, GREEN, 10)
line("도서명  정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션 (한영대역)", 9, True)
line("‘정본역(正本譯)’은 공인본문(Textus Receptus)과 맛소라 본문을 저본으로 삼았음을 뜻하는 말이며, 다른 번역본의 가치를 부정하는 표현이 아닙니다.")
line("이 책의 한국어 본문은 킹제임스 성경(KJV, 1611)의 영어 본문과 그 저본인 공인본문(Textus Receptus)·맛소라 본문을 히브리어·아람어·헬라어 원문과 대조하여 바이블 인사이트가 직접 번역한 것입니다. 기존 한국어 역본을 저본으로 삼지 않은 독자적인 번역이며, 번역 원칙 전문은 biblynote.com/translation 에 공개되어 있습니다.")
line("절마다 왼쪽에 한국어 본문, 오른쪽에 King James Version 영어 본문을 나란히 배치하여 한 절씩 대조하며 읽을 수 있습니다.")
line("본문 소제목은 독자의 이해를 돕기 위하여 바이블 인사이트가 새로 지은 것으로, 성경 원문의 일부가 아닙니다.")
line("영어 본문(King James Version, 1611)은 대한민국 저작권법상 보호 기간이 만료된 퍼블릭 도메인 저작물입니다.")
line("본문 서체는 SIL Open Font License로 배포되는 Noto Serif CJK KR을 사용하였습니다.")
line("한국어 번역 저작권 ⓒ 오광일 · 바이블 인사이트, 2026. 이 책의 한국어 본문을 출판사의 서면 허락 없이 복제·전재·배포할 수 없습니다. 다만 개인 묵상·설교·강의·논문에서의 통상적인 인용은 출처(정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션)를 밝히는 조건으로 허용합니다.")
line("펴낸곳  바이블 인사이트 출판사", 9, False, None, 8)
line("옮긴이  오광일")
line("문의  contact@biblynote.com · biblynote.com")

# ── 본문 ──
body = doc.add_section(WD_SECTION.NEW_PAGE)
body.page_width, body.page_height = Cm(15.2), Cm(22.5)
body.top_margin = Cm(1.6); body.bottom_margin = Cm(1.4)
body.left_margin = body.right_margin = Cm(1.2)
add_page_number(body)

HALF = Cm(6.4)

def add_seg_table(rows_data):
    """구간(소제목 사이) 절들을 2열 표로 — 좌 한글 / 우 영어"""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.autofit = False
    set_insideV(tbl)
    for (n, ko, en), row in zip(rows_data, tbl.rows):
        c1, c2 = row.cells
        c1.width = HALF; c2.width = HALF
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3.5); p1.paragraph_format.line_spacing = 1.4
        kf(p1.add_run(str(n) + " "), 6.5, True, GREEN)
        kf(p1.add_run(ko))
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(3.5); p2.paragraph_format.line_spacing = 1.35
        kf(p2.add_run(str(n) + " "), 6.5, True, GREEN)
        kf(p2.add_run(en), 8.2)

t0 = time.time(); total_v = 0
for bi, bk in enumerate(books):
    if bi > 0:
        doc.add_page_break()
    center(bk["ko"], 19, True, INK, 8, 2)
    center(bk["en"].upper(), 9.5, False, GRAY, 0, 9)
    for ch in range(1, bk["ch"]+1):
        vs = en_ch(bk["file"], ch)
        if vs is None:
            continue
        hd = headings_for(bk["file"], ch)
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(8); hp.paragraph_format.space_after = Pt(4)
        hp.paragraph_format.keep_with_next = True
        kf(hp.add_run(str(ch)), 20, True, GREEN)
        kf(hp.add_run("  %s · %s %d" % (bk["ko"], bk["en"], ch)), 8, False, GRAY)
        seg = []
        for n, v in enumerate(vs):
            if (n+1) in hd:
                if seg:
                    add_seg_table(seg); seg = []
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(7); sp.paragraph_format.space_after = Pt(3)
                sp.paragraph_format.keep_with_next = True
                kf(sp.add_run(hd[n+1]), 9.6, True, RGBColor(0x4a, 0x5a, 0x52))
            seg.append((n+1, clean_ko(v.get("ko", "") or ""), (v.get("en", "") or "").strip()))
            total_v += 1
        if seg:
            add_seg_table(seg)
    if (bi+1) % 10 == 0:
        print("  [%d/66] %s %.0fs" % (bi+1, bk["ko"], time.time()-t0), flush=True)

doc.save(OUT)
print("저장:", OUT, "| 절 %s | %.0fs |" % (format(total_v, ","), time.time()-t0),
      round(os.path.getsize(OUT)/1e6, 1), "MB")
