# -*- coding: utf-8 -*-
"""정본역 킹제임스 성경 — 앞·뒤 부록 공용 데이터 모듈 (2026-08-05 구성 확정)

  앞부분: ①간행사 ②번역 원칙 6개조 ③저본 명세·정본역 정의 ④일러두기 ⑤약자표
  뒷부분: A 교리 용어 해설 · B 죽음과 심판 원어 도표 · D 구원으로 가는 길 ·
          G 도량형·화폐 환산표 · L QR 안내
  판별 특화: J 암송 구절 30선(큰글자판) · K KJV 고어 사전(한영대역)

  각 섹션 = {"id","title","blocks":[block...]}
  block   = {"t":"p"|"h"|"lead"|"table"|"verse"|"kv"|"note", ...}
    p     본문 문단 {"text"}
    h     소제목    {"text"}
    lead  강조 문단 {"text"}
    table 표        {"head":[...], "rows":[[...],...]}
    verse 성구 인용 {"ref","text"}  (text는 빌드 시 bible/kr에서 동적 주입)
    kv    용어 항목 {"term","desc"}
    note  작은 안내 {"text"}
  사용: front_sections() / back_sections(edition)  edition∈{"bigprint","parallel"}
"""
import json, io, os, re

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_books = json.load(io.open(os.path.join(ROOT, "bible", "books.json"), encoding="utf-8"))

_KO_QUOTES = re.compile(u'[“”‘’"\'「」『』《》〈〉]')
def _clean(s):
    return re.sub(r"\s{2,}", " ", _KO_QUOTES.sub("", str(s))).strip()

_KO2FILE = {b["ko"]: b["file"] for b in _books}
def verse_text(ref_ko):
    """'요한복음 3:16' → 정본역 본문(따옴표 정리 적용)"""
    m = re.match(r"(.+?)\s+(\d+):(\d+)$", ref_ko.strip())
    book, ch, v = m.group(1), int(m.group(2)), int(m.group(3))
    p = os.path.join(ROOT, "bible", "kr", "%s-%d.json" % (_KO2FILE[book], ch))
    vs = json.loads(io.open(p, encoding="utf-8-sig").read())
    return _clean(vs[v - 1])

# 표준 약자(개역 관례를 따르되 저작권 무관한 관용 약칭)
_ABBR = ["창","출","레","민","신","수","삿","룻","삼상","삼하","왕상","왕하","대상","대하","스","느","에",
         "욥","시","잠","전","아","사","렘","애","겔","단","호","욜","암","옵","욘","미","나","합","습","학","슥","말",
         "마","막","눅","요","행","롬","고전","고후","갈","엡","빌","골","살전","살후","딤전","딤후","딛","몬",
         "히","약","벧전","벧후","요일","요이","요삼","유","계"]

QR_URL = "https://biblynote.com/bible"
QR_PNG = os.path.join(ROOT, "책원고", "출판준비", "_qr_bible.png")

def make_qr():
    import qrcode
    q = qrcode.QRCode(border=2, box_size=10)
    q.add_data(QR_URL); q.make(fit=True)
    img = q.make_image(fill_color="#232d28", back_color="white")
    os.makedirs(os.path.dirname(QR_PNG), exist_ok=True)
    img.save(QR_PNG)
    return QR_PNG

# ═══════════════════ 앞부분 ═══════════════════

def front_sections():
    S = []
    S.append({"id": "preface", "title": "펴내며", "blocks": [
        {"t": "p", "text": "성경은 하나님의 말씀입니다. 그러므로 성경을 옮기는 일은 한 글자도 가볍게 다룰 수 없는 일이었습니다."},
        {"t": "p", "text": "이 성경은 킹제임스 성경(KJV, 1611)과 그 저본인 공인본문(Textus Receptus)·맛소라 본문을 히브리어·아람어·헬라어 원문과 한 절씩 대조하며, 다섯 해에 걸쳐 새로 옮긴 것입니다. 기존 한국어 역본을 저본으로 삼지 않았고, 핵심 교리 용어는 한국 교회가 백여 년 지켜 온 전통 표기를 그대로 보존하였습니다. 새롭게 하되 바꾸지 말아야 할 것은 바꾸지 않는 것 — 그것이 이 번역이 지킨 첫 원칙입니다."},
        {"t": "p", "text": "성경을 읽는 사람이 누구의 해석에도 기대지 않고 말씀 앞에 바로 서는 것, 그리하여 스스로 읽고 분별하는 성도가 세워지는 것이 이 책의 기도입니다."},
        {"t": "p", "text": "말씀으로 시대를 읽는 모든 분들에게 이 성경을 드립니다."},
        {"t": "note", "text": "바이블 인사이트 출판사 · 옮긴이 오광일"},
    ]})
    S.append({"id": "principles", "title": "번역 원칙", "blocks": [
        {"t": "kv", "term": "1. 공인본문의 온전한 보존",
         "desc": "공인본문(Textus Receptus)에 있는 어구와 절은 하나도 덜어내지 않고 전부 옮겼습니다. 현대 비평 본문이 생략하는 구절들(마태복음 17:21, 사도행전 8:37, 요한일서 5:7 등)도 본문 그대로 보존하였습니다."},
        {"t": "kv", "term": "2. 원문이 최종 근거",
         "desc": "KJV 영어 본문을 출발점으로 삼되, 핵심 단어는 히브리어·아람어·헬라어 원문을 직접 확인하여 원문이 최종 근거가 되게 하였습니다."},
        {"t": "kv", "term": "3. 죽음과 심판 용어의 구분",
         "desc": "원어가 구분하는 것은 번역도 구분합니다. 스올·하데스는 '음부'로, 게헨나는 '지옥'으로 옮겨 죽은 자의 세계와 최후의 형벌 장소를 섞지 않았습니다."},
        {"t": "kv", "term": "4. 하나님의 이름",
         "desc": "구약의 신명사문자(יהוה)는 전권에서 '여호와'로 통일하였습니다. 신약에는 '여호와'가 한 번도 등장하지 않습니다 — 신약 원문에 신명사문자가 없기 때문입니다."},
        {"t": "kv", "term": "5. 핵심 교리 용어의 전통 보존",
         "desc": "독생자·칭의·성화·구속·중생·영생·침례 등 교리의 뼈대가 되는 용어는 새 말을 만들지 않고 한국 교회의 전통 표기를 지켰습니다."},
        {"t": "kv", "term": "6. 문체와 문장기호",
         "desc": "본문은 권위 있는 성경 문어체로 통일하고, 따옴표류 문장기호는 한글 성경의 관례를 따라 쓰지 않았습니다. 문장의 구성과 어순은 어느 역본도 베끼지 않은 독자적 번역입니다."},
        {"t": "note", "text": "번역 원칙 전문과 절별 번역 해설은 biblynote.com/translation 에 공개되어 있습니다."},
    ]})
    S.append({"id": "sources", "title": "저본(底本)", "blocks": [
        {"t": "lead", "text": "'정본역(正本譯)'은 공인본문(Textus Receptus)과 맛소라 본문을 저본으로 삼았음을 뜻하는 이름이며, 다른 번역본의 가치를 부정하는 표현이 아닙니다."},
        {"t": "table", "head": ["구분", "저본", "비고"],
         "rows": [["구약", "맛소라 본문 (Masoretic Text)", "히브리어·아람어"],
                  ["신약", "공인본문 (Textus Receptus)", "헬라어"],
                  ["대조", "킹제임스 성경 (KJV, 1611)", "영어 — 대한민국 저작권법상 퍼블릭 도메인"]]},
        {"t": "p", "text": "본문 구성은 KJV와 동일한 66권 1,189장 31,102절이며, 절 구분도 1:1로 일치합니다."},
    ]})
    S.append({"id": "guide", "title": "일러두기", "blocks": [
        {"t": "kv", "term": "소제목", "desc": "본문 중의 소제목은 독자의 이해를 돕기 위하여 바이블 인사이트가 새로 지은 것으로, 성경 원문의 일부가 아닙니다."},
        {"t": "kv", "term": "절 번호", "desc": "각 절 앞의 작은 숫자는 절 번호입니다. 장 시작의 큰 숫자는 장 번호입니다."},
        {"t": "kv", "term": "각주", "desc": "본문 이해에 필요한 원어 설명과 번역 해설은 biblynote.com 성경 읽기에서 절별로 볼 수 있습니다."},
        {"t": "kv", "term": "음부와 지옥", "desc": "'음부'(스올·하데스)는 죽은 자들이 머무는 곳을, '지옥'(게헨나)은 최후의 불 형벌을 가리킵니다. 원어가 다르므로 번역도 구분하였습니다."},
        {"t": "kv", "term": "왕국", "desc": "'하나님의 왕국'과 '하늘의 왕국'(마태복음 고유 표현)은 원문의 구분을 따라 나누어 옮겼습니다."},
        {"t": "kv", "term": "문장기호", "desc": "대화문에도 따옴표를 쓰지 않는 한글 성경의 관례를 따랐습니다."},
    ]})
    # 약자표
    rows, half = [], (len(_books) + 1) // 2
    for i in range(half):
        L = _books[i]; R = _books[i + half] if i + half < len(_books) else None
        rows.append([L["ko"], _ABBR[i], (R["ko"] if R else ""), (_ABBR[i + half] if R else "")])
    S.append({"id": "abbr", "title": "성경 각 권과 약자", "blocks": [
        {"t": "table", "head": ["책명", "약자", "책명", "약자"], "rows": rows},
    ]})
    return S

# ═══════════════════ 뒷부분 ═══════════════════

_TERMS = [
    ("독생자(獨生子)", "요 3:16 — '오직 하나뿐인 아들'(모노게네스). 예수 그리스도께서 아버지와 본질을 같이하시는 유일한 아들이심을 나타내는 말."),
    ("칭의(稱義)", "하나님께서 믿는 자를 의롭다고 선언하시는 것(롬 3:24). 행위의 결과가 아니라 그리스도의 의가 전가되는 법정적 선언."),
    ("성화(聖化)", "의롭다 하심을 받은 자가 성령으로 점점 거룩하게 되어 가는 과정(살전 4:3)."),
    ("구속(救贖)", "값을 지불하고 종을 사서 자유하게 함. 그리스도께서 자기 피로 우리를 죄에서 사신 일(엡 1:7)."),
    ("중생(重生)·거듭남", "위로부터 새로 나는 것(요 3:3). 성령으로 말미암은 새 생명의 시작."),
    ("영생(永生)", "단지 끝없이 사는 것이 아니라, 하나님을 아는 생명(요 17:3). 믿는 자에게 지금 주어지는 하나님의 생명."),
    ("침례(浸禮)", "밥티조 — '물에 잠그다'. 그리스도와 함께 죽고 함께 살아남을 나타내는 순종의 첫걸음(롬 6:4)."),
    ("회개(悔改)", "메타노이아 — 마음과 방향을 돌이킴. 단순한 후회가 아니라 하나님께로 돌아서는 것(행 3:19)."),
    ("화목제물", "힐라스모스 — 하나님의 공의로운 진노를 만족시키는 제물. 그리스도께서 우리 죄를 위한 화목제물이 되심(요일 2:2)."),
    ("언약(言約)", "하나님께서 맺으시고 하나님께서 지키시는 약속. 옛 언약(율법)과 새 언약(그리스도의 피, 눅 22:20)."),
    ("은혜(恩惠)", "받을 자격이 없는 자에게 값없이 주시는 하나님의 호의(엡 2:8)."),
    ("믿음", "바라는 것들의 실상이요 보이지 않는 것들의 증거(히 11:1). 구원은 오직 믿음으로 받는다."),
    ("의(義)", "하나님의 기준에 맞는 올바름. 사람의 의는 누더기 같으나(사 64:6) 하나님의 의가 믿는 자에게 주어진다."),
    ("보혜사(保惠師)", "파라클레토스 — 곁에 불러 돕는 이. 성령님(요 14:16)과 대언자 되시는 그리스도(요일 2:1)."),
    ("성령(聖靈)", "삼위일체 하나님의 제3위. 거듭나게 하시고, 내주하시며, 인치시고, 거룩하게 하시는 하나님의 영."),
    ("삼위일체", "성부·성자·성령 세 위격이 한 하나님이심. '이 셋은 하나이니라'(요일 5:7)."),
    ("음부(陰府)", "스올(히)·하데스(헬) — 죽은 자들이 머무는 곳. 최후 심판 전의 중간 상태(눅 16:23)."),
    ("지옥(地獄)", "게헨나 — 최후의 불 형벌 장소. 몸과 혼을 능히 지옥에 멸하시는 이를 두려워하라(마 10:28)."),
    ("부활(復活)", "몸이 다시 사는 것. 그리스도께서 첫 열매로 살아나셨고, 믿는 자도 그와 같이 살아난다(고전 15:20-23)."),
    ("휴거(携擧)", "주께서 강림하실 때 살아 남은 성도가 공중으로 끌어올려져 주를 영접하는 일(살전 4:17)."),
    ("재림(再臨)", "그리스도께서 다시 오심. 공중 강림(성도를 위해)과 지상 재림(성도와 함께)으로 성취된다."),
    ("천년왕국", "그리스도께서 재림하여 땅 위에 세우실 천 년의 나라(계 20:4-6). 이스라엘에 약속된 메시아 왕국의 성취."),
    ("첫째 부활", "생명의 부활. 첫째 부활에 참여하는 자는 둘째 사망이 다스리지 못한다(계 20:6)."),
    ("둘째 사망", "크고 흰 보좌 심판 뒤 불못에 던져지는 최종 상태(계 20:14). 생명책에 기록된 자에게는 해당이 없다."),
    ("생명책", "어린양의 생명책. 구원받은 모든 자의 이름이 기록된 책(계 21:27)."),
    ("적그리스도", "그리스도를 대적하며 그리스도를 자칭하는 자. 마지막 때 나타날 불법의 사람(살후 2:3-4)."),
    ("대환난", "마지막 때 온 땅에 임할 7년의 환난(마 24:21). 야곱의 고난의 때(렘 30:7)."),
    ("복음(福音)", "유앙겔리온 — 좋은 소식. 그리스도께서 성경대로 죽으시고 장사되시고 사흘 만에 살아나신 것(고전 15:3-4)."),
    ("속죄(贖罪)", "죄를 덮음(히 카파르). 피 흘림이 없이는 죄 사함이 없다(히 9:22)."),
    ("전가(轉嫁)", "그리스도의 의가 믿는 자의 것으로 옮겨져 계산됨. 아브라함이 하나님을 믿으매 그것이 의로 여겨졌다(롬 4:3)."),
    ("성막(聖幕)·성전", "하나님께서 사람 가운데 거하시는 처소. 그리스도의 몸(요 2:21)과 성도의 몸(고전 6:19)의 예표."),
    ("유월절", "어린양의 피로 죽음이 넘어간 밤(출 12장). 우리의 유월절 어린양 되신 그리스도의 예표(고전 5:7)."),
    ("안식일", "일곱째 날의 쉼. 참 안식은 그리스도 안에서 완성된다(히 4:9-10)."),
    ("할례(割禮)", "언약의 표. 신약에서는 손으로 하지 아니한 마음의 할례로 성취된다(골 2:11)."),
    ("선지자(先知者)", "하나님의 말씀을 맡아 대신 전하는 자. 그 직분의 중심은 예고가 아니라 말씀 선포다."),
    ("사도(使徒)", "보냄을 받은 자. 부활하신 주님을 목격하고 교회의 기초를 놓은 열두 사도와 바울."),
    ("교회(敎會)", "에클레시아 — 불러냄을 받은 자들의 모임. 건물이 아니라 그리스도를 머리로 하는 몸(엡 1:22-23)."),
    ("장로·감독·목자", "한 직분의 세 이름(행 20:17,28). 무리를 치는 자가 아니라 양 무리의 본이 되는 자(벧전 5:3)."),
    ("경건(敬虔)", "하나님을 향한 삶의 자세. 경건의 능력은 부인하고 모양만 있는 시대를 경계하라(딤후 3:5)."),
    ("세상", "코스모스 — 하나님을 떠난 질서 전체. 세상을 사랑하지 말라(요일 2:15), 그러나 하나님은 세상을 사랑하사 아들을 주셨다(요 3:16)."),
    ("육(肉)과 영(靈)", "타락한 본성(육)과 거듭난 생명(영)의 대립. 육으로 난 것은 육이요 영으로 난 것은 영이니라(요 3:6)."),
    ("혼(魂)과 영(靈)", "사람은 영과 혼과 몸으로 지음받았다(살전 5:23). 말씀은 혼과 영을 찔러 쪼갠다(히 4:12)."),
]

_ARCHAIC = [
    ("thee / thou / thy / thine", "너를 / 네가 / 너의 / 너의 것 — 2인칭 단수. KJV는 단수(thou)와 복수(ye/you)를 구분한다"),
    ("ye / you", "너희가 / 너희를 — 2인칭 복수"),
    ("hath / hast", "has / have(2인칭 단수) — ~을 가지고 있다"),
    ("doth / dost", "does / do(2인칭 단수)"),
    ("saith", "says — 말하다"), ("spake", "spoke — 말하였다"),
    ("shalt / wilt", "shall / will의 2인칭 단수 — ~하리라"),
    ("art", "are의 2인칭 단수 — thou art 네가 ~이다"),
    ("unto", "to — ~에게, ~까지"), ("thereof", "of it — 그것의"),
    ("wherefore", "therefore / why — 그러므로, 어찌하여"),
    ("verily", "truly — 진실로"), ("behold", "보라"),
    ("begat", "낳았다(족보)"), ("begotten", "낳은 — only begotten 독생하신"),
    ("henceforth", "이제부터"), ("hitherto", "지금까지"),
    ("whence", "어디로부터"), ("thither / hither", "그리로 / 이리로"),
    ("wax", "점점 ~하게 되다 — waxed old 낡아지다"),
    ("suffer", "허락하다(고전 의미) — Suffer the little children 어린아이들을 용납하라"),
    ("charity", "사랑(아가페) — 고전 13장의 charity는 자선이 아니라 사랑"),
    ("conversation", "행실, 생활 방식(고전 의미)"),
    ("meat", "음식(고기만이 아님)"), ("corn", "곡식(옥수수가 아님)"),
    ("kine", "암소들"), ("fowl", "새, 날짐승"),
    ("raiment", "의복"), ("victuals", "양식"),
    ("morrow", "다음 날 — on the morrow 이튿날"),
    ("even / eventide", "저녁"), ("sup", "저녁을 먹다"),
    ("anon / straightway", "곧, 즉시"), ("by and by", "즉시(현대어와 반대 의미)"),
    ("peradventure", "혹시, 아마"), ("haply", "혹시"),
    ("nay / yea", "아니라 / 그러하다"), ("nigh", "가까이"),
    ("betwixt", "between — 사이에"), ("amongst", "among — 가운데"),
    ("whosoever", "누구든지"), ("whatsoever", "무엇이든지"),
    ("thereon / therein", "그 위에 / 그 안에"),
    ("hearken", "귀 기울여 듣다"), ("beseech", "간구하다"),
    ("quicken", "살리다 — quickened 살림을 받은"),
    ("let", "막다(고전 의미도 있음 — 살후 2:7 he who now letteth 막는 자)"),
    ("prevent", "앞서 가다(고전 의미 — 살전 4:15 shall not prevent 앞서지 못하리라)"),
    ("ghost", "영 — Holy Ghost 성령, gave up the ghost 숨을 거두다"),
    ("mete", "재다 — with what measure ye mete 너희가 헤아리는 그 헤아림으로"),
    ("ere", "~하기 전에"), ("oft / oftentimes", "자주"),
    ("sore", "심히 — sore afraid 심히 두려워하다"),
    ("twain", "둘 — they twain shall be one flesh 둘이 한 몸이 되리라"),
    ("fain", "기꺼이"), ("list", "원하다 — the wind bloweth where it listeth 바람이 임의로 불매"),
    ("bewray", "드러내다"), ("holden", "붙잡힌"),
    ("shew", "show — 보이다"), ("wot / wist", "알다 / 알았다"),
    ("would God", "~라면 좋으련만(간절한 소원)"),
    ("God forbid", "결코 그럴 수 없느니라(메 게노이토)"),
]

def back_sections(edition="bigprint"):
    S = []
    S.append({"id": "terms", "title": "부록 1 · 핵심 교리 용어 해설", "blocks":
        [{"t": "p", "text": "이 성경이 지켜 쓴 전통 교리 용어들의 뜻을 간추렸습니다. 더 자세한 풀이는 biblynote.com 성경사전에서 볼 수 있습니다."}]
        + [{"t": "kv", "term": t, "desc": d} for t, d in _TERMS]})
    S.append({"id": "death", "title": "부록 2 · 죽음과 심판 — 원어 구분 도표", "blocks": [
        {"t": "p", "text": "여러 역본이 '지옥' 한 단어로 뭉뚱그린 곳을, 이 성경은 원어를 따라 구분하여 옮겼습니다."},
        {"t": "table", "head": ["원어", "번역", "뜻", "대표 구절"], "rows": [
            ["שְׁאוֹל 스올 (구약)", "음부", "죽은 자들이 머무는 곳", "창 37:35 · 시 16:10"],
            ["ᾅδης 하데스 (신약)", "음부", "죽은 자들의 세계 — 부자와 나사로", "눅 16:23 · 계 20:13"],
            ["γέεννα 게헨나", "지옥", "최후의 불 형벌 장소", "마 10:28 · 막 9:43"],
            ["ταρταρόω 타르타로스", "지옥(타르타로스)", "범죄한 천사들을 가둔 곳", "벧후 2:4"]]},
        {"t": "p", "text": "음부는 최후 심판 전의 중간 상태이고, 지옥(게헨나)과 불못은 최후의 형벌입니다. 사망과 음부도 마지막에는 불못에 던져집니다(계 20:14)."},
    ]})
    S.append({"id": "salvation", "title": "부록 3 · 구원으로 가는 길", "blocks": [
        {"t": "lead", "text": "성경이 말하는 구원의 길은 어렵지 않습니다. 다음 다섯 걸음을 말씀 그대로 따라가 보십시오."},
        {"t": "h", "text": "1. 모든 사람이 죄인입니다"},
        {"t": "verse", "ref": "로마서 3:23"},
        {"t": "h", "text": "2. 죄의 값은 사망입니다"},
        {"t": "verse", "ref": "로마서 6:23"},
        {"t": "h", "text": "3. 그리스도께서 나를 위해 죽으셨습니다"},
        {"t": "verse", "ref": "로마서 5:8"},
        {"t": "h", "text": "4. 입으로 시인하고 마음으로 믿으면"},
        {"t": "verse", "ref": "로마서 10:9"},
        {"t": "h", "text": "5. 누구든지 주의 이름을 부르는 자는"},
        {"t": "verse", "ref": "로마서 10:13"},
        {"t": "p", "text": "지금 이 자리에서, 자신의 말로 기도하십시오. 하나님은 마음의 중심을 보십니다. 구원은 행위가 아니라 믿음으로, 값없이 주시는 하나님의 선물입니다(에베소서 2:8-9)."},
    ]})
    S.append({"id": "measures", "title": "부록 4 · 도량형·화폐 환산표", "blocks": [
        {"t": "h", "text": "길이"},
        {"t": "table", "head": ["단위", "환산", "비고"], "rows": [
            ["손가락 너비", "약 1.9cm", "렘 52:21"],
            ["손바닥 너비", "약 7.5cm", "네 손가락"],
            ["뼘", "약 22.5cm", "손을 편 길이"],
            ["규빗", "약 45cm", "팔꿈치에서 손끝까지"],
            ["갈대(측량 자)", "약 2.7m", "6규빗 — 겔 40:5"],
            ["스다디온", "약 185m", "계 21:16"],
            ["안식일에 가기 알맞은 길", "약 1km", "행 1:12"]]},
        {"t": "h", "text": "무게·화폐"},
        {"t": "table", "head": ["단위", "환산", "비고"], "rows": [
            ["게라", "약 0.57g", "세겔의 1/20"],
            ["세겔", "약 11.4g", "은 세겔 = 노동자 나흘 품삯"],
            ["미나(마네)", "약 570g", "50세겔"],
            ["달란트", "약 34kg", "3,000세겔 — 품꾼 약 20년 품삯"],
            ["데나리온", "은전 하나", "품꾼 하루 품삯(마 20:2)"],
            ["렙돈", "가장 작은 동전", "과부의 두 렙돈(막 12:42)"],
            ["앗사리온", "동전 하나", "참새 두 마리 값(마 10:29)"]]},
        {"t": "h", "text": "부피"},
        {"t": "table", "head": ["단위", "환산", "비고"], "rows": [
            ["오멜", "약 2.3L", "만나 한 사람 하루 분량(출 16:16)"],
            ["에바", "약 23L", "10오멜"],
            ["호멜", "약 230L", "10에바"],
            ["밧", "약 23L", "액체 단위"],
            ["힌", "약 3.8L", "출 29:40"]]},
    ]})
    if edition == "bigprint":
        mem = json.load(io.open(os.path.join(ROOT, "content", "memorize.json"), encoding="utf-8"))
        pick_ids = ["jn3-16","ro3-23","ro6-23","ro5-8","ro10-9","eph2-8","jn1-12","jn5-24","jn14-6","ac4-12",
                    "ro8-1","ro8-28","php4-6","php4-13","php4-19","ps23-1","ps119-105","pr3-5","isa41-10","isa53-5",
                    "mt6-33","mt11-28","jos1-9","jer29-11","2co5-17","gal2-20","heb11-1","jas1-5","1jn1-9","rev3-20"]
        by_id = {v["id"]: v for v in mem["verses"]}
        vv = [by_id[i] for i in pick_ids if i in by_id]
        if len(vv) < 30:                                   # id 매칭 부족 시 앞에서 보충
            seen = {v["id"] for v in vv}
            vv += [v for v in mem["verses"] if v["id"] not in seen][:30 - len(vv)]
        S.append({"id": "memorize", "title": "부록 5 · 암송하면 좋은 말씀 30선", "blocks":
            [{"t": "p", "text": "마음에 새기면 평생 힘이 되는 말씀들입니다. 하루에 한 절씩, 소리 내어 읽으며 외워 보십시오."}]
            + [{"t": "verse", "ref": v["ref"], "text": _clean(v["kr"])} for v in vv[:30]]})
    if edition == "parallel":
        S.append({"id": "archaic", "title": "부록 5 · KJV 고어(古語) 작은 사전", "blocks":
            [{"t": "p", "text": "킹제임스 성경(1611)의 영어에는 오늘날 쓰지 않는 옛말이 많습니다. 자주 나오는 것들을 간추렸습니다. 뜻을 알고 읽으면 KJV 영문이 한결 가깝게 다가옵니다."}]
            + [{"t": "kv", "term": t, "desc": d} for t, d in _ARCHAIC]})
    S.append({"id": "online", "title": "부록 %d · 온라인에서 더 깊이" % (6 if edition in ("bigprint", "parallel") else 5), "blocks": [
        {"t": "p", "text": "이 성경의 전체 본문을 절별 원어 각주와 함께 온라인에서 읽을 수 있습니다. 성경사전 5,000여 항목, 강의 320여 편, 매일 묵상이 함께 제공됩니다."},
        {"t": "qr", "url": QR_URL, "caption": "biblynote.com — 바이블 인사이트"},
    ]})
    return S

def resolve_verses(sections):
    """verse 블록에 정본역 본문 주입(이미 text가 있으면 유지)"""
    for s in sections:
        for b in s["blocks"]:
            if b["t"] == "verse" and not b.get("text"):
                b["text"] = verse_text(b["ref"])
    return sections

# ═══════════════════ 공용 렌더러 ═══════════════════
# 5개 빌더(PDF 2·Word 2·EPUB 2)가 아래 렌더러만 호출한다.

def pdf_flowables(sections, col_w, big=False):
    """reportlab 플로어블 목록 — 섹션마다 새 페이지 + 러닝헤드 마커(_bibly_head)"""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.platypus import Image as RLImage
    GREEN = HexColor("#00593c"); INK = HexColor("#232d28"); GRAY = HexColor("#8a8a8a")
    RULE = HexColor("#c9c2b4"); SUB = HexColor("#4a5a52")
    bs = 11.5 if big else 9.2                       # 본문 크기(판형별)
    s_t  = ParagraphStyle("axt",  fontName="NSKB", fontSize=bs+7, leading=(bs+7)*1.4,
                          textColor=INK, spaceBefore=6, spaceAfter=10)
    s_h  = ParagraphStyle("axh",  fontName="NSKB", fontSize=bs+1.5, leading=(bs+1.5)*1.45,
                          textColor=GREEN, spaceBefore=9, spaceAfter=4)
    s_p  = ParagraphStyle("axp",  fontName="NSK",  fontSize=bs, leading=bs*1.72,
                          alignment=TA_JUSTIFY, textColor=INK, spaceAfter=5, wordWrap="CJK")
    s_ld = ParagraphStyle("axl",  parent=s_p, fontName="NSKB", textColor=SUB,
                          spaceBefore=3, spaceAfter=8)
    s_kv = ParagraphStyle("axk",  parent=s_p, spaceAfter=6)
    s_v  = ParagraphStyle("axv",  parent=s_p, leftIndent=4*mm, spaceAfter=7)
    s_nt = ParagraphStyle("axn",  fontName="NSK", fontSize=bs-1.2, leading=(bs-1.2)*1.6,
                          textColor=GRAY, spaceBefore=6)
    s_tb = ParagraphStyle("axtb", fontName="NSK", fontSize=bs-0.8, leading=(bs-0.8)*1.5,
                          textColor=INK, wordWrap="CJK")
    s_th = ParagraphStyle("axth", parent=s_tb, fontName="NSKB", textColor=GREEN)
    s_ct = ParagraphStyle("axc",  parent=s_p, alignment=TA_CENTER)
    def E(x):
        return str(x).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    out = []
    for sec in sections:
        out.append(PageBreak())
        tp = Paragraph(E(sec["title"]), s_t)
        tp._bibly_head = (sec["title"].split(" · ")[-1], "APPENDIX" if "부록" in sec["title"] else "")
        out.append(tp)
        for b in sec["blocks"]:
            t = b["t"]
            if t == "p":
                out.append(Paragraph(E(b["text"]), s_p))
            elif t == "h":
                out.append(Paragraph(E(b["text"]), s_h))
            elif t == "lead":
                out.append(Paragraph(E(b["text"]), s_ld))
            elif t == "note":
                out.append(Paragraph(E(b["text"]), s_nt))
            elif t == "kv":
                out.append(Paragraph('<font name="NSKB">%s</font>  %s' % (E(b["term"]), E(b["desc"])), s_kv))
            elif t == "verse":
                out.append(Paragraph('<font name="NSKB" color="#00593c">%s</font>  %s'
                                     % (E(b["ref"]), E(b.get("text", ""))), s_v))
            elif t == "table":
                ncol = len(b["head"])
                data = [[Paragraph(E(h), s_th) for h in b["head"]]] + \
                       [[Paragraph(E(c), s_tb) for c in row] for row in b["rows"]]
                if ncol == 3:                       # 첫 열(단위·용어)을 살짝 넓게
                    w0 = col_w * 0.34
                    ws = [w0] + [(col_w - w0) / 2] * 2
                else:
                    ws = [col_w / ncol] * ncol
                tb = Table(data, colWidths=ws, repeatRows=1)
                tb.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.8, GREEN),
                    ("LINEBELOW", (0, 1), (-1, -1), 0.3, RULE),
                    ("TOPPADDING", (0, 0), (-1, -1), 2.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
                out += [Spacer(1, 2), tb, Spacer(1, 5)]
            elif t == "qr":
                make_qr()
                out += [Spacer(1, 8), RLImage(QR_PNG, width=30*mm, height=30*mm, hAlign="CENTER"),
                        Spacer(1, 3), Paragraph(E(b.get("caption", b["url"])), s_ct)]
    return out


def docx_render(doc, sections, font, base=10.5, first_break=True):
    """python-docx 렌더 — 호출부의 현재 섹션(1단 권장)에 이어서 그린다.
       표는 내장 Table Grid 스타일만 사용(tblPr 수동 조작 금지 — 과거 손상 사고 예방)."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    GREEN = RGBColor(0x00, 0x59, 0x3c); INK = RGBColor(0x23, 0x2d, 0x28)
    GRAY = RGBColor(0x8a, 0x8a, 0x8a); SUB = RGBColor(0x4a, 0x5a, 0x52)
    def kf(run, size=None, bold=False, color=None):
        run.font.name = font
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
        if size: run.font.size = Pt(size)
        run.bold = bold
        if color: run.font.color.rgb = color
        return run
    def para(after=4, before=0, line=1.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.line_spacing = line
        return p
    for si, sec in enumerate(sections):
        if si > 0 or first_break:
            doc.add_page_break()
        tp = para(after=10, before=4)
        kf(tp.add_run(sec["title"]), base + 7, True, INK)
        for b in sec["blocks"]:
            t = b["t"]
            if t == "p":
                kf(para().add_run(b["text"]), base)
            elif t == "h":
                p = para(after=3, before=8); p.paragraph_format.keep_with_next = True
                kf(p.add_run(b["text"]), base + 1.5, True, GREEN)
            elif t == "lead":
                kf(para(after=8, before=2).add_run(b["text"]), base, True, SUB)
            elif t == "note":
                kf(para(before=6).add_run(b["text"]), base - 1.5, False, GRAY)
            elif t == "kv":
                p = para(after=5)
                kf(p.add_run(b["term"] + "  "), base, True, INK)
                kf(p.add_run(b["desc"]), base)
            elif t == "verse":
                p = para(after=6); p.paragraph_format.left_indent = Cm(0.35)
                kf(p.add_run(b["ref"] + "  "), base, True, GREEN)
                kf(p.add_run(b.get("text", "")), base)
            elif t == "table":
                tbl = doc.add_table(rows=1, cols=len(b["head"]))
                tbl.style = "Table Grid"
                for h, cell in zip(b["head"], tbl.rows[0].cells):
                    kf(cell.paragraphs[0].add_run(h), base - 1, True, GREEN)
                for row in b["rows"]:
                    cells = tbl.add_row().cells
                    for c, cell in zip(row, cells):
                        cp = cell.paragraphs[0]
                        cp.paragraph_format.space_after = Pt(1)
                        kf(cp.add_run(str(c)), base - 1)
                para(after=6)
            elif t == "qr":
                make_qr()
                p = para(before=8); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(QR_PNG, width=Cm(3.2))
                p2 = para(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                kf(p2.add_run(b.get("caption", b["url"])), base - 0.5)


EPUB_CSS = """
h1.apx{font-size:1.35em;text-align:left;margin:1.1em 0 0.8em;color:#232d28;}
p.lead{font-weight:bold;color:#4a5a52;margin:0.4em 0 0.9em;}
p.kv{margin:0 0 0.55em;}
p.apxv{margin:0 0 0.6em 0.7em;}
table.apx{border-collapse:collapse;width:100%;margin:0.5em 0 1em;font-size:0.92em;}
table.apx th{border-bottom:2px solid #00593c;color:#00593c;text-align:left;padding:0.25em 0.4em;}
table.apx td{border-bottom:1px solid #d9d3c6;padding:0.25em 0.4em;vertical-align:top;}
"""

def epub_xhtml_sections(sections, qr_src="images/qr.png"):
    """섹션별 (id, title, body_html) 목록 — 빌더가 자기 xhtml() 래퍼로 감싼다"""
    import html as _h
    E = _h.escape
    out = []
    for sec in sections:
        parts = ['<h1 class="apx">%s</h1>' % E(sec["title"])]
        for b in sec["blocks"]:
            t = b["t"]
            if t == "p":
                parts.append("<p>%s</p>" % E(b["text"]))
            elif t == "h":
                parts.append("<h2>%s</h2>" % E(b["text"]))
            elif t == "lead":
                parts.append('<p class="lead">%s</p>' % E(b["text"]))
            elif t == "note":
                parts.append('<p class="small">%s</p>' % E(b["text"]))
            elif t == "kv":
                parts.append('<p class="kv"><b>%s</b>  %s</p>' % (E(b["term"]), E(b["desc"])))
            elif t == "verse":
                parts.append('<p class="apxv"><b style="color:#00593c;">%s</b>  %s</p>'
                             % (E(b["ref"]), E(b.get("text", ""))))
            elif t == "table":
                rows = ["<tr>%s</tr>" % "".join("<th>%s</th>" % E(h) for h in b["head"])]
                rows += ["<tr>%s</tr>" % "".join("<td>%s</td>" % E(c) for c in row) for row in b["rows"]]
                parts.append('<table class="apx">%s</table>' % "".join(rows))
            elif t == "qr":
                make_qr()
                parts.append('<p class="center"><img src="%s" alt="QR" style="width:9em;max-width:60%%;"/><br/>%s</p>'
                             % (qr_src, E(b.get("caption", b["url"]))))
        out.append((sec["id"], sec["title"], "".join(parts)))
    return out


if __name__ == "__main__":
    import sys
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    fs = resolve_verses(front_sections())
    for ed in ("bigprint", "parallel"):
        bs = resolve_verses(back_sections(ed))
        n = sum(len(s["blocks"]) for s in bs)
        print("%s: 앞 %d섹션 · 뒤 %d섹션(%d블록)" % (ed, len(fs), len(bs), n))
    make_qr()
    print("QR 생성:", QR_PNG)
