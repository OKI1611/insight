# -*- coding: utf-8 -*-
"""정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션 · 큰글자판 — Word(.docx) 생성

  PDF 정본(tools/build_pdfs.py)과 같은 기준(2026-08-03 확정):
    2단 배열(자이언트 프린트형) · 큰 숫자 장 표기 · 소제목 수록 ·
    한글 따옴표류 제거 · '한글 전용' 표기 없음 · 쪽번호(바닥글)
  ※ 표지·판권은 1단 섹션, 본문은 2단 섹션.
  실행: python tools/build_bigprint_docx.py
"""
import json, io, os, re, sys, time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import build_appendix as APX                       # 앞·뒤 부록(2026-08-05 확정 구성)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT  = os.path.join(ROOT, "책원고", "정본역킹제임스성경_큰글자판.docx")
books = json.load(io.open(os.path.join(ROOT, "bible", "books.json"), encoding="utf-8"))

try:
    HEADINGS = json.load(io.open(os.path.join(ROOT, "bible", "headings", "ko.json"), encoding="utf-8"))
except Exception:
    HEADINGS = {}
def headings_for(book, ch):
    return {v: t for v, t in HEADINGS.get(book, {}).get(str(ch), [])}

_KO_QUOTES = re.compile(u'[“”‘’"\'「」『』《》〈〉]')
def clean_ko(s):
    return re.sub(r"\s{2,}", " ", _KO_QUOTES.sub("", str(s))).strip()

def kr(b, c):
    return json.loads(io.open(os.path.join(ROOT, "bible", "kr", "%s-%d.json" % (b, c)), encoding="utf-8-sig").read())

FONT  = "Noto Serif CJK KR"                    # SIL OFL — 유료 판매 임베딩 안전
GREEN = RGBColor(0x00, 0x59, 0x3c)
GRAY  = RGBColor(0x8a, 0x8a, 0x8a)
INK   = RGBColor(0x23, 0x2d, 0x28)
SUBG  = RGBColor(0x00, 0x59, 0x3c)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)          # A4
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(1.8)
sec.left_margin = sec.right_margin = Cm(1.9)

st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(14)
st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(4.5)
st.paragraph_format.line_spacing = 1.45

def kf(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if size: run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color
    return run

def center(txt, size, bold=False, color=None, before=0, after=6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    kf(p.add_run(txt), size, bold, color)
    return p

def set_columns(section, n, space_cm=0.9):
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); section._sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(int(space_cm * 567)))       # twips
    cols.set(qn("w:sep"), "1" if n > 1 else "0")            # 단 구분선

def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "17"); rPr.append(sz)  # 8.5pt
    r.append(rPr); t = OxmlElement("w:t"); t.text = "1"; r.append(t)
    fld.append(r); p._p.append(fld)

# ── 표지 (1단 섹션) ──
for _ in range(6): doc.add_paragraph()
# 역본명이 27자라 한 줄에 안 들어감 → 2줄로 나눠 배치
center("정본역(正本譯) 킹제임스 성경", 30, True, INK, 0, 2)
center("헤리티지 에디션", 20, True, INK)
center("큰글자판", 24, True, GREEN)
center("전 66권", 14, False, GRAY)
for _ in range(9): doc.add_paragraph()
center("바이블 인사이트 출판사", 13, True)

# ── 판권 ──
doc.add_page_break()
for _ in range(2): doc.add_paragraph()
def line(txt, size=11, bold=False, color=None, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    kf(p.add_run(txt), size, bold, color)
line("일러두기 · 판권", 15, True, GREEN, 12)
line("도서명  정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션 (큰글자판)", 11, True)
line("‘정본역(正本譯)’은 공인본문(Textus Receptus)과 맛소라 본문을 저본으로 삼았음을 뜻하는 말이며, 다른 번역본의 가치를 부정하는 표현이 아닙니다.")
line("이 책의 한국어 본문은 킹제임스 성경(KJV, 1611)의 영어 본문과 그 저본인 공인본문(Textus Receptus)·맛소라 본문을 히브리어·아람어·헬라어 원문과 대조하여 바이블 인사이트가 직접 번역한 것입니다. 기존 한국어 역본을 저본으로 삼지 않은 독자적인 번역이며, 번역 원칙 전문은 biblynote.com/translation 에 공개되어 있습니다.")
line("본문 소제목은 독자의 이해를 돕기 위하여 바이블 인사이트가 새로 지은 것으로, 성경 원문의 일부가 아닙니다.")
line("눈이 편안하도록 본문 글자를 크게 하고 줄 간격을 넉넉하게 조판하였습니다.")
line("본문 서체는 SIL Open Font License로 배포되는 Noto Serif CJK KR을 사용하였습니다.")
line("한국어 번역 저작권 ⓒ 오광일 · 바이블 인사이트, 2026. 이 책의 한국어 본문을 출판사의 서면 허락 없이 복제·전재·배포할 수 없습니다. 다만 개인 묵상·설교·강의·논문에서의 통상적인 인용은 출처(정본역(正本譯) 킹제임스 성경 : 헤리티지 에디션)를 밝히는 조건으로 허용합니다.")
line("펴낸곳  바이블 인사이트 출판사", 11, False, None, 10)
line("옮긴이  오광일")
line("문의  contact@biblynote.com · biblynote.com")

# ── 앞부록 (1단 유지 — 표 조판에 유리) ──
APX.docx_render(doc, APX.resolve_verses(APX.front_sections()), FONT, base=12)

# ── 본문 (2단 섹션) ──
body = doc.add_section(WD_SECTION.NEW_PAGE)
body.page_width, body.page_height = Cm(21), Cm(29.7)
body.top_margin = Cm(2.2); body.bottom_margin = Cm(1.8)
body.left_margin = body.right_margin = Cm(1.9)
set_columns(body, 2)
add_page_number(body)

t0 = time.time()
for bi, bk in enumerate(books):
    if bi > 0:
        doc.add_page_break()
    center(bk["ko"], 25, True, INK, 10, 2)
    center(bk["en"].upper(), 11, False, GRAY, 0, 12)
    for ch in range(1, bk["ch"]+1):
        vs = kr(bk["file"], ch)
        hd = headings_for(bk["file"], ch)
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10); hp.paragraph_format.space_after = Pt(5)
        hp.paragraph_format.keep_with_next = True
        kf(hp.add_run(str(ch)), 30, True, GREEN)
        kf(hp.add_run("  " + bk["ko"]), 11, False, GRAY)
        for n, v in enumerate(vs):
            if not isinstance(v, str):
                continue
            if (n+1) in hd:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(11); sp.paragraph_format.space_after = Pt(4)
                sp.paragraph_format.keep_with_next = True
                kf(sp.add_run(hd[n+1]), 13.5, True, SUBG)
            p = doc.add_paragraph()
            kf(p.add_run(str(n+1) + " "), 9.5, True, GREEN)
            kf(p.add_run(clean_ko(v)))
    if (bi+1) % 10 == 0:
        print("  [%d/66] %s %.0fs" % (bi+1, bk["ko"], time.time()-t0), flush=True)

# ── 뒤부록 (1단 새 섹션) ──
apx = doc.add_section(WD_SECTION.NEW_PAGE)
apx.page_width, apx.page_height = Cm(21), Cm(29.7)
apx.top_margin = Cm(2.2); apx.bottom_margin = Cm(1.8)
apx.left_margin = apx.right_margin = Cm(1.9)
set_columns(apx, 1)
add_page_number(apx)
APX.docx_render(doc, APX.resolve_verses(APX.back_sections("bigprint")), FONT, base=12, first_break=False)

doc.save(OUT)
print("저장:", OUT, "| %.0fs |" % (time.time()-t0), round(os.path.getsize(OUT)/1e6, 1), "MB")
